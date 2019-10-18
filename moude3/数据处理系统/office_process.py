# -*- coding=utf-8 -*-
#author:wang

from docx import Document
import openpyxl

class WordOperation:
    """
    word操作类，可以根据不同的输入内容，生成word文件，要求至少完成三类内容的输入，标题，副标题，正文
    """
    def __init__(self, title, subtitle, text):
        self.title = title
        self.subtitle = subtitle
        self.text = text
    def write_word(self):
        "写入word并保存"
        try:
            doc = Document()
            doc.add_heading(self.title)
            doc.add_paragraph(self.subtitle,'Subtitle')
            doc.add_paragraph(self.text)
            name = self.title + '.docx'
            doc.save(name)
            print('word文件保存完成！')
        except Exception as f:
            print(f)
    
class ExcelOperation:
    '''
    根据第一列的数据，按年进行拆分，放到新的工作表，例：2015年数据，表名为2015
    按年拆分后的数据，在数据最后一行，添加平均价格
    from_file_name:数据文件名
    sheet_name:数据表名
    go_file_name:数据新建表名
    '''
    def __init__(self, from_file_name, sheet_name, go_file_name):
        self.wb = openpyxl.load_workbook(from_file_name)
        self.sh1 = self.wb[sheet_name]
        self.go_name = go_file_name

    def run(self):
        "完成逻辑控制"
        l = []
        for rows in self.sh1.rows:
            if rows[0].coordinate != 'A1':
                l.append(rows[0].value[:4])
                l = list(set(l))

        for x in l:
            self.write_data(x)

        self.wb.save(self.go_name)
        print('excel保存完成')


    def write_data(self, sheet_name):
        """
        写入表格
        sheet_name :数据保存的表名
        """
        index = 2
        self.wb.create_sheet(sheet_name)
        sh2 = self.wb[sheet_name]
        sh2['A1'] = self.sh1['A1'].value
        sh2['B1'] = self.sh1['B1'].value
        for rows in self.sh1.rows:
            if rows[0].coordinate != 'A1' and rows[0].value[:4] == sheet_name:
                sh2['A' + str(index)] = rows[0].value
                sh2['B' + str(index)] = rows[1].value
                index += 1
        num = sh2.max_row
        # 求平均值
        l=[]
        for row in sh2.rows:
            if row[0].coordinate != 'A1':
                l.append(int(row[1].value))
        average_num = sum(l)/(num-1)
        
        sh2.cell(row=num+1, column=2).value = average_num
        sh2.cell(row=num+1, column=1).value = '平均分'

class OfficeOperation:
    "选择需要的操作类，逻辑选择"
    def operation(self):
        while True:
            print('请选择需要进行的操作')
            menu = {
                '1':'操作word',
                '2':'操作excel',            
                '0':'quit'
            }
            for k,v in menu.items():
                print(k, v)
            try:
                choose = input('请输入需要选择的操作编号：')           
                if choose == '0':
                    print('退出！' )
                    break
                elif choose == '1':
                    WO = WordOperation('wangshuai', 'python', 'text')
                    WO.write_word()
                elif choose == '2':
                    EO = ExcelOperation('btc.xlsx', 'btc', 'btc-1.xlsx')
                    EO.run()
            except Exception as f:
                print("输入有误请重新输入")

def main():
    while True:
        print('请选择需要进行的操作')
        menu = {
            '1':'操作word',
            '2':'操作excel',            
            '0':'quit'
        }
        for k,v in menu.items():
            print(k, v)
        try:
            choose = input('请输入需要选择的操作编号：')           
            if choose == '0':
                print('退出！' )
                break
            elif choose == '1':
                WO = WordOperation('wangshuai', 'python', 'text')
                WO.write_word()
            elif choose == '2':
                EO = ExcelOperation('btc.xlsx', 'btc', 'btc-1.xlsx')
                EO.run()
        except Exception as f:
            print("输入有误请重新输入")
        
if __name__ == "__main__":
    main()